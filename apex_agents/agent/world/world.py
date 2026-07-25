"""On-disk world environment for the APEX-Agents agent.

Each APEX-Agents task runs against a "world" — a snapshot of a
professional knowledge worker's file surface (an investment-banking
deal room: spreadsheets, PDFs, memos). The world's assets ship as a
``world_files_zipped/{world_id}.zip`` on the HF dataset; optional
per-task input files live under ``task_files/{task_id}/**``.

:class:`WorldFiles` extracts the world zip + the task files into a
temporary directory and exposes a small read-only file surface the
agent's domain tools sit on top of:

* :meth:`list_files` — recursive relative-path listing.
* :meth:`read_text` — UTF-8 text read (with a byte cap).
* :meth:`read_spreadsheet` — xlsx cell dump via ``openpyxl``.
* :meth:`read_pdf` — extracted PDF text via ``pypdf``.
* :meth:`search` — case-insensitive substring grep over text files.

Heavy parsers (``openpyxl`` / ``pypdf``) are imported lazily so the
module stays importable offline; tests use :class:`FakeWorld` and
never touch this class.
"""

from __future__ import annotations

import logging
import os
import shutil
import tempfile
import zipfile
from collections.abc import Callable, Mapping
from pathlib import Path
from typing import Any

from ...data.dataset import ApexAgentsRecord


logger = logging.getLogger(__name__)


# Per-case world factory: ``(record) -> WorldFiles``-like.
WorldFactory = Callable[[ApexAgentsRecord], "WorldFiles"]


__all__ = ["WorldFactory", "WorldFiles", "build_world_factory", "world_factory_from_mapping"]


# Cap on a single text/file read so a hostile or huge asset can't blow
# the agent's context window or the host's memory.
_MAX_READ_BYTES = 200_000
_MAX_SEARCH_FILE_BYTES = 1_000_000


class WorldFiles:
    """Read-only view over an extracted APEX-Agents world directory."""

    def __init__(self, root: str | os.PathLike[str], *, owns_root: bool = False) -> None:
        self._root = Path(root).resolve()
        self._owns_root = owns_root

    @property
    def root(self) -> Path:
        return self._root

    # ─── construction ─────────────────────────────────────────────────

    @classmethod
    def from_zip(
        cls,
        world_zip_path: str | os.PathLike[str],
        *,
        task_files_dir: str | os.PathLike[str] | None = None,
        extract_dir: str | os.PathLike[str] | None = None,
    ) -> "WorldFiles":
        """Extract a world zip (+ optional task files) into a temp dir.

        ``world_files_zipped/{world_id}.zip`` contains a ``filesystem/``
        and ``.apps_data/`` subtree. Per-task input files (if any) are
        copied in over the top so a task can override / add world
        assets.
        """
        owns_root = extract_dir is None
        target = Path(extract_dir) if extract_dir is not None else Path(tempfile.mkdtemp(prefix="apex_world_"))
        target.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(world_zip_path) as zf:
            zf.extractall(target)
        if task_files_dir is not None:
            src = Path(task_files_dir)
            if src.is_dir():
                shutil.copytree(src, target, dirs_exist_ok=True)
        return cls(target, owns_root=owns_root)

    def close(self) -> None:
        """Remove the extracted tree if this instance owns it."""
        if self._owns_root and self._root.exists():
            shutil.rmtree(self._root, ignore_errors=True)

    # ─── path safety ──────────────────────────────────────────────────

    def _resolve(self, rel_path: str) -> Path:
        """Resolve a caller path inside the world, rejecting traversal."""
        candidate = (self._root / rel_path).resolve()
        if candidate != self._root and self._root not in candidate.parents:
            raise ValueError(f"Path {rel_path!r} escapes the world root.")
        return candidate

    # ─── read surface ─────────────────────────────────────────────────

    def list_files(self, subdir: str = "") -> list[str]:
        """Return the sorted list of files (relative paths) under ``subdir``."""
        base = self._resolve(subdir) if subdir else self._root
        if not base.exists():
            return []
        out: list[str] = []
        for path in sorted(base.rglob("*")):
            if path.is_file():
                out.append(str(path.relative_to(self._root)))
        return out

    def read_text(self, rel_path: str, *, max_bytes: int = _MAX_READ_BYTES) -> str:
        """Read a UTF-8 text file (errors replaced), capped at ``max_bytes``."""
        path = self._resolve(rel_path)
        if not path.is_file():
            raise FileNotFoundError(rel_path)
        data = path.read_bytes()[:max_bytes]
        return data.decode("utf-8", errors="replace")

    def read_spreadsheet(
        self,
        rel_path: str,
        *,
        sheet: str | None = None,
        max_rows: int = 200,
        max_cols: int = 40,
    ) -> str:
        """Read an ``.xlsx`` workbook, one sheet at a time.

        A financial model is a multi-tab workbook: the first tab is
        usually a narrative cover ("Executive Summary"); the computed
        outputs live on later tabs. Concatenating every sheet into one
        blob means downstream observation truncation hides the model
        tabs — the agent only ever sees the cover. So:

        * The reply ALWAYS begins with a ``# Sheets:`` index (every
          sheet name + its row x col dimensions). It is small and
          survives truncation, so the agent can always see what tabs
          exist and navigate to the right one.
        * ``sheet=<name>`` renders that sheet's cells (tab-separated,
          capped). Without ``sheet``, only the index + the FIRST sheet
          are returned, with an explicit hint to re-call with
          ``sheet=`` for the others.
        """
        path = self._resolve(rel_path)
        if not path.is_file():
            raise FileNotFoundError(rel_path)
        # Legacy .xls (BIFF) — openpyxl is .xlsx-only; route to xlrd.
        if path.suffix.lower() == ".xls":
            return self._read_xls(path, sheet=sheet, max_rows=max_rows, max_cols=max_cols)
        try:
            from openpyxl import load_workbook
        except ImportError as exc:  # pragma: no cover - import guard
            raise ImportError("read_spreadsheet requires the optional `openpyxl` dependency.") from exc
        workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        try:
            names = list(workbook.sheetnames)
            index = " | ".join(
                f"{nm} ({workbook[nm].max_row or '?'}x{workbook[nm].max_column or '?'})" for nm in names
            )
            header = "# Sheets: " + index

            if sheet is not None and sheet not in names:
                return (
                    f"{header}\nERROR: no sheet named {sheet!r}. Re-call read_spreadsheet with "
                    f"one of the exact sheet names listed above via the `sheet` argument."
                )

            target = sheet if sheet is not None else (names[0] if names else None)
            if target is None:
                return header
            chunks: list[str] = [header]
            if sheet is None and len(names) > 1:
                chunks.append(
                    f"(Showing only the first sheet {target!r}. The computed model is usually on "
                    f"a LATER tab — re-call read_spreadsheet with sheet='<exact name above>'.)"
                )
            chunks.append(f"# Sheet: {target}")
            for row_index, row in enumerate(workbook[target].iter_rows(values_only=True)):
                if row_index >= max_rows:
                    chunks.append("... (truncated rows)")
                    break
                cells = ["" if c is None else str(c) for c in row[:max_cols]]
                chunks.append("\t".join(cells))
            return "\n".join(chunks)
        finally:
            workbook.close()

    def _read_xls(self, path: Path, *, sheet: str | None, max_rows: int, max_cols: int) -> str:
        """Render a legacy ``.xls`` workbook with the same shape as the xlsx path."""
        try:
            import xlrd
        except ImportError as exc:  # pragma: no cover - import guard
            raise ImportError("reading .xls requires the optional `xlrd` dependency.") from exc
        book = xlrd.open_workbook(str(path))
        names = book.sheet_names()
        index = " | ".join(f"{nm} ({book.sheet_by_name(nm).nrows}x{book.sheet_by_name(nm).ncols})" for nm in names)
        header = "# Sheets: " + index
        if sheet is not None and sheet not in names:
            return (
                f"{header}\nERROR: no sheet named {sheet!r}. Re-call read_spreadsheet with "
                f"one of the exact sheet names listed above via the `sheet` argument."
            )
        target = sheet if sheet is not None else (names[0] if names else None)
        if target is None:
            return header
        chunks: list[str] = [header]
        if sheet is None and len(names) > 1:
            chunks.append(
                f"(Showing only the first sheet {target!r}. Re-call read_spreadsheet with "
                f"sheet='<exact name above>' for the others.)"
            )
        chunks.append(f"# Sheet: {target}")
        ws = book.sheet_by_name(target)
        for r in range(min(ws.nrows, max_rows)):
            cells = ["" if v is None else str(v) for v in ws.row_values(r)[:max_cols]]
            chunks.append("\t".join(cells))
        if ws.nrows > max_rows:
            chunks.append("... (truncated rows)")
        return "\n".join(chunks)

    def read_pdf(self, rel_path: str, *, max_pages: int = 50) -> str:
        """Extract text from a PDF (best-effort) via ``pypdf``.

        Several APEX world files carry a ``.pdf`` name but are actually
        ZIP/Office containers (xlsx/docx) — pypdf fails opaquely on
        those. We sniff the magic bytes and return an actionable
        message routing the agent to the right reader instead.
        """
        path = self._resolve(rel_path)
        if not path.is_file():
            raise FileNotFoundError(rel_path)
        head = path.read_bytes()[:8]
        if head[:4] == b"PK\x03\x04":
            return (
                f"{rel_path!r} is NOT a PDF — it is a ZIP/Office container (xlsx or docx) "
                "mislabeled .pdf. Re-read it with read_spreadsheet (if a workbook) or "
                "read_docx (if a Word document)."
            )
        if head[:5] != b"%PDF-":
            return f"{rel_path!r} does not look like a PDF (magic={head[:5]!r}). Try read_file/read_docx/read_spreadsheet."
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - import guard
            raise ImportError("read_pdf requires the optional `pypdf` dependency.") from exc
        try:
            reader = PdfReader(str(path))
        except Exception as exc:  # pragma: no cover - defensive
            return f"Could not parse {rel_path!r} as a PDF ({type(exc).__name__}: {exc}). Try read_file."
        parts: list[str] = []
        for page_index, page in enumerate(reader.pages):
            if page_index >= max_pages:
                parts.append("... (truncated pages)")
                break
            try:
                parts.append(page.extract_text() or "")
            except Exception:  # pragma: no cover - defensive
                parts.append("")
        return "\n".join(parts)

    def read_docx(self, rel_path: str, *, max_bytes: int = _MAX_READ_BYTES) -> str:
        """Extract text from a ``.docx`` (Word) document via ``python-docx``.

        Legal worlds are .docx-heavy (contracts, surveys). Without this
        the agent's only option was read_file, which returns binary
        gibberish for .docx — making those cases unwinnable. Joins
        paragraph text and table cell text, byte-capped.
        """
        path = self._resolve(rel_path)
        if not path.is_file():
            raise FileNotFoundError(rel_path)
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - import guard
            raise ImportError("read_docx requires the optional `python-docx` dependency.") from exc
        try:
            doc = Document(str(path))
        except Exception as exc:  # pragma: no cover - defensive
            return f"Could not parse {rel_path!r} as a .docx ({type(exc).__name__}: {exc})."
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for trow in table.rows:
                cells = [c.text.strip() for c in trow.cells]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)[:max_bytes]

    def search(self, query: str, *, max_results: int = 50) -> list[dict[str, Any]]:
        """Case-insensitive substring grep over the world's text files.

        Returns ``[{"file": rel, "line": n, "text": line}, ...]``.
        Binary / oversized files are skipped.
        """
        needle = query.casefold()
        if not needle:
            return []
        hits: list[dict[str, Any]] = []
        for rel in self.list_files():
            path = self._root / rel
            try:
                if path.stat().st_size > _MAX_SEARCH_FILE_BYTES:
                    continue
                text = path.read_text(encoding="utf-8", errors="strict")
            except (UnicodeDecodeError, OSError):
                continue
            for line_no, line in enumerate(text.splitlines(), start=1):
                if needle in line.casefold():
                    hits.append({"file": rel, "line": line_no, "text": line.strip()[:300]})
                    if len(hits) >= max_results:
                        return hits
        return hits


def build_world_factory(
    *,
    cache_dir: str | None = None,
    repo_id: str = "mercor/apex-agents",
) -> WorldFactory:
    """Return a ``(record) -> WorldFiles`` factory backed by HuggingFace.

    Downloads ``world_files_zipped/{world_id}.zip`` and (when present)
    ``task_files/{task_id}/`` from the HF dataset, extracts them into a
    temp dir, and yields a :class:`WorldFiles`. Imported lazily +
    pushed behind a factory so production runs hit HF while tests
    inject a :class:`FakeWorld` factory and never download anything.
    """

    def _factory(record: ApexAgentsRecord) -> WorldFiles:
        from huggingface_hub import hf_hub_download, snapshot_download

        world_id = record.world_id
        task_id = record.task_id
        world_zip = hf_hub_download(
            repo_id=repo_id,
            filename=f"world_files_zipped/{world_id}.zip",
            repo_type="dataset",
            cache_dir=cache_dir,
        )
        task_files_dir: str | None = None
        try:
            snapshot_root = snapshot_download(
                repo_id=repo_id,
                repo_type="dataset",
                allow_patterns=[f"task_files/{task_id}/*"],
                cache_dir=cache_dir,
            )
            candidate = Path(snapshot_root) / "task_files" / task_id
            if candidate.is_dir():
                task_files_dir = str(candidate)
        except Exception:  # pragma: no cover - optional per-task files
            logger.debug("No per-task input files for %s", task_id, exc_info=True)
        return WorldFiles.from_zip(world_zip, task_files_dir=task_files_dir)

    return _factory


def world_factory_from_mapping(mapping: Mapping[str, WorldFiles]) -> WorldFactory:
    """Build a ``(record) -> WorldFiles`` factory from a prebuilt mapping.

    Convenience for callers that already extracted worlds (keyed by
    ``world_id``); avoids re-downloading per case.
    """

    def _factory(record: ApexAgentsRecord) -> WorldFiles:
        world_id = record.world_id
        if world_id not in mapping:
            raise KeyError(f"No prebuilt world for world_id={world_id!r}.")
        return mapping[world_id]

    return _factory
