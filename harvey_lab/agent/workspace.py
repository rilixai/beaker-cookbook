"""Per-task on-disk workspace for the Harvey LAB agent.

Each LAB task ships a ``documents/`` folder (the source record) and a set
of named ``deliverables`` the agent must produce. :class:`TaskWorkspace`
lays that out on disk exactly like Harvey's harness — a read-only
``documents/`` subtree plus a writable ``output/`` subtree — and exposes
the small file surface the Stirrup tools sit on top of:

* :meth:`list_files` — recursive listing of ``documents/`` + ``output/``.
* :meth:`read_document` — text extraction for .docx / .xlsx / .pdf / .eml
  / plain text, byte-capped.
* :meth:`search_documents` — case-insensitive grep over the documents.
* :meth:`write_deliverable` / :meth:`edit_deliverable` — create / patch an
  ``output/`` file (relative paths are always rooted in ``output/``).
* :meth:`collect_deliverables` — read every ``output/`` file back out for
  the rubric judge.

The document parsers (python-docx / openpyxl / pypdf) are imported lazily
so the module stays importable offline; hermetic tests build workspaces
from a fixture directory and never touch the heavy parsers unless a test
provides a real binary.
"""

from __future__ import annotations

import email
import logging
import shutil
from collections.abc import Callable, Mapping
from pathlib import Path
from typing import Any


logger = logging.getLogger(__name__)


# Per-case factory: ``(record) -> TaskWorkspace``. Production fetches the task
# documents from the pinned upstream commit; tests inject a fixture-backed one.
TaskSource = Callable[[Any], "TaskWorkspace"]


__all__ = [
    "TaskSource",
    "TaskWorkspace",
    "build_bundled_task_source",
    "build_github_task_source",
    "task_source_from_dir",
    "task_source_from_mapping",
]


class TaskWorkspace:
    """A materialized ``documents/`` + ``output/`` workspace for one task."""

    def __init__(self, root: str | Path, *, max_document_chars: int = 200_000) -> None:
        self._root = Path(root).resolve()
        self._documents = self._root / "documents"
        self._output = self._root / "output"
        self._documents.mkdir(parents=True, exist_ok=True)
        self._output.mkdir(parents=True, exist_ok=True)
        self._max_document_chars = max_document_chars

    @property
    def root(self) -> Path:
        return self._root

    @property
    def documents_dir(self) -> Path:
        return self._documents

    @property
    def output_dir(self) -> Path:
        return self._output

    def close(self) -> None:
        """Remove the workspace tree."""
        shutil.rmtree(self._root, ignore_errors=True)

    # ─── path safety ──────────────────────────────────────────────────

    def _resolve_within(self, base: Path, rel_path: str) -> Path:
        candidate = (base / rel_path).resolve()
        if candidate != base and base not in candidate.parents:
            raise ValueError(f"Path {rel_path!r} escapes {base.name}/.")
        return candidate

    # ─── read surface (documents/) ────────────────────────────────────

    def list_files(self) -> list[str]:
        """Return sorted ``documents/…`` + ``output/…`` relative paths."""
        out: list[str] = []
        for base in (self._documents, self._output):
            for path in sorted(base.rglob("*")):
                if path.is_file():
                    out.append(str(path.relative_to(self._root)))
        return out

    def read_document(self, rel_path: str) -> str:
        """Extract text from a document under ``documents/`` (or ``output/``).

        A leading ``documents/`` or ``output/`` component selects the subtree;
        a bare name defaults to ``documents/``.
        """
        path = self._route_read(rel_path)
        if not path.is_file():
            raise FileNotFoundError(rel_path)
        suffix = path.suffix.lower()
        if suffix == ".docx":
            text = _read_docx(path)
        elif suffix in (".xlsx", ".xlsm"):
            text = _read_xlsx(path)
        elif suffix == ".pdf":
            text = _read_pdf(path)
        elif suffix == ".eml":
            text = _read_eml(path)
        else:
            text = path.read_bytes().decode("utf-8", errors="replace")
        return text[: self._max_document_chars]

    def _route_read(self, rel_path: str) -> Path:
        parts = Path(rel_path).parts
        if parts and parts[0] == "output":
            return self._resolve_within(self._output, str(Path(*parts[1:])))
        if parts and parts[0] == "documents":
            return self._resolve_within(self._documents, str(Path(*parts[1:])))
        return self._resolve_within(self._documents, rel_path)

    def search_documents(self, query: str, *, max_results: int = 50) -> list[dict[str, Any]]:
        """Case-insensitive substring grep over readable document text."""
        needle = query.casefold()
        if not needle:
            return []
        hits: list[dict[str, Any]] = []
        for path in sorted(self._documents.rglob("*")):
            if not path.is_file():
                continue
            rel = str(path.relative_to(self._root))
            try:
                text = self.read_document(rel)
            except Exception:  # noqa: BLE001 - unreadable file, skip
                continue
            for line_no, line in enumerate(text.splitlines(), start=1):
                if needle in line.casefold():
                    hits.append({"file": rel, "line": line_no, "text": line.strip()[:300]})
                    if len(hits) >= max_results:
                        return hits
        return hits

    # ─── write surface (output/) ──────────────────────────────────────

    def write_deliverable(self, rel_path: str, content: str) -> str:
        """Write ``content`` to ``output/<rel_path>`` (parents created)."""
        path = self._resolve_within(self._output, rel_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(content, encoding="utf-8")
        return str(path.relative_to(self._root))

    def edit_deliverable(self, rel_path: str, old: str, new: str) -> str:
        """Replace the first occurrence of ``old`` with ``new`` in an output file."""
        path = self._resolve_within(self._output, rel_path)
        if not path.is_file():
            raise FileNotFoundError(f"output/{rel_path}")
        text = path.read_text(encoding="utf-8")
        if old not in text:
            raise ValueError(f"edit_deliverable: snippet not found in output/{rel_path}.")
        path.write_text(text.replace(old, new, 1), encoding="utf-8")
        return str(path.relative_to(self._root))

    def collect_deliverables(self) -> dict[str, str]:
        """Read every file under ``output/`` back out, keyed by filename."""
        out: dict[str, str] = {}
        for path in sorted(self._output.rglob("*")):
            if path.is_file():
                rel = str(path.relative_to(self._output))
                out[rel] = path.read_bytes().decode("utf-8", errors="replace")
        return out


# ─── lazy document parsers ────────────────────────────────────────────


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - import guard
        raise ImportError("read_document requires the optional `python-docx` dependency for .docx.") from exc
    try:
        doc = Document(str(path))
    except Exception as exc:  # pragma: no cover - defensive
        return f"Could not parse {path.name!r} as .docx ({type(exc).__name__}: {exc})."
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _read_xlsx(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - import guard
        raise ImportError("read_document requires the optional `openpyxl` dependency for .xlsx.") from exc
    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    try:
        names = list(workbook.sheetnames)
        chunks: list[str] = ["# Sheets: " + " | ".join(names)]
        for name in names:
            chunks.append(f"# Sheet: {name}")
            for row_index, row in enumerate(workbook[name].iter_rows(values_only=True)):
                if row_index >= 200:
                    chunks.append("... (truncated rows)")
                    break
                chunks.append("\t".join("" if c is None else str(c) for c in row[:40]))
        return "\n".join(chunks)
    finally:
        workbook.close()


def _read_pdf(path: Path) -> str:
    head = path.read_bytes()[:5]
    if head[:4] == b"PK\x03\x04":
        return f"{path.name!r} is a ZIP/Office container mislabeled .pdf; read it as .docx/.xlsx instead."
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - import guard
        raise ImportError("read_document requires the optional `pypdf` dependency for .pdf.") from exc
    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # pragma: no cover - defensive
        return f"Could not parse {path.name!r} as a PDF ({type(exc).__name__}: {exc})."
    parts: list[str] = []
    for page_index, page in enumerate(reader.pages):
        if page_index >= 50:
            parts.append("... (truncated pages)")
            break
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - defensive per-page
            parts.append("")
    return "\n".join(parts)


def _read_eml(path: Path) -> str:
    msg = email.message_from_bytes(path.read_bytes())
    headers = [f"{h}: {msg.get(h, '')}" for h in ("From", "To", "Cc", "Date", "Subject") if msg.get(h)]
    body_parts: list[str] = []
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                payload = part.get_payload(decode=True)
                if isinstance(payload, bytes):
                    body_parts.append(payload.decode("utf-8", errors="replace"))
    else:
        payload = msg.get_payload(decode=True)
        if isinstance(payload, bytes):
            body_parts.append(payload.decode("utf-8", errors="replace"))
    return "\n".join(headers) + "\n\n" + "\n".join(body_parts)


# ─── task-source factories ────────────────────────────────────────────


def task_source_from_dir(tasks_root: str | Path, *, max_document_chars: int = 200_000) -> TaskSource:
    """Build a ``(record) -> TaskWorkspace`` factory backed by a local tree.

    ``tasks_root`` is a directory laid out like ``harveyai/harvey-labs``:
    ``<practice_area>/<slug>/documents/*``. The record's ``task_id`` is the
    ``<practice_area>/<slug>`` path. Used by the export script and by tests
    (which point it at a fixture tree) — fully offline.
    """
    base = Path(tasks_root)

    def _factory(record: Any) -> TaskWorkspace:
        import tempfile

        task_id = str(getattr(record, "task_id", "") or "")
        src_docs = base / task_id / "documents"
        ws = TaskWorkspace(tempfile.mkdtemp(prefix="harvey_lab_"), max_document_chars=max_document_chars)
        if src_docs.is_dir():
            shutil.copytree(src_docs, ws.documents_dir, dirs_exist_ok=True)
        return ws

    return _factory


def build_github_task_source(
    *,
    repo: str,
    commit: str,
    max_document_chars: int = 200_000,
) -> TaskSource:
    """Build a factory that fetches a task's documents from a pinned commit.

    Mirrors the apex recipe's HF world factory: production runs pull the
    per-task ``documents/`` from ``raw.githubusercontent.com/<repo>/<commit>/
    tasks/<task_id>/documents/<file>`` (the LAB benchmark is a public repo),
    materialize them into a temp :class:`TaskWorkspace`, and return it. The
    document filenames come from ``record.documents`` (recorded at export time),
    so no directory-listing API call is needed. Imported lazily + behind a
    factory so tests never hit the network.
    """

    def _factory(record: Any) -> TaskWorkspace:
        import tempfile

        task_id = str(getattr(record, "task_id", "") or "")
        documents: tuple[str, ...] = tuple(getattr(record, "documents", ()) or ())
        ws = TaskWorkspace(tempfile.mkdtemp(prefix="harvey_lab_"), max_document_chars=max_document_chars)
        for name in documents:
            url = f"https://raw.githubusercontent.com/{repo}/{commit}/tasks/{task_id}/documents/{name}"
            dest = ws.documents_dir / name
            dest.parent.mkdir(parents=True, exist_ok=True)
            dest.write_bytes(_fetch_bytes(url))
        return ws

    return _factory


# raw.githubusercontent.com rate-limits unauthenticated requests; many cases
# fetch concurrently, so a bare urlopen hits HTTP 429. Retry transient failures
# (429 + 5xx + connection errors) with exponential backoff honouring
# ``Retry-After`` when present.
_FETCH_MAX_ATTEMPTS = 6
_FETCH_BASE_DELAY = 2.0
_FETCH_MAX_DELAY = 60.0


def _fetch_bytes(url: str) -> bytes:
    import random
    import time
    import urllib.error
    import urllib.request

    last_exc: Exception | None = None
    for attempt in range(_FETCH_MAX_ATTEMPTS):
        req = urllib.request.Request(url, headers={"User-Agent": "harvey-lab-cookbook/1.0"})
        try:
            with urllib.request.urlopen(req, timeout=60) as resp:  # noqa: S310 - pinned https host
                return bytes(resp.read())
        except urllib.error.HTTPError as exc:
            if exc.code != 429 and exc.code < 500:
                raise
            last_exc = exc
            retry_after = exc.headers.get("Retry-After") if exc.headers else None
            delay = float(retry_after) if retry_after and retry_after.isdigit() else None
        except urllib.error.URLError as exc:
            last_exc = exc
            delay = None
        if attempt == _FETCH_MAX_ATTEMPTS - 1:
            break
        if delay is None:
            delay = min(_FETCH_BASE_DELAY * (2**attempt), _FETCH_MAX_DELAY)
        time.sleep(delay + random.uniform(0, 1.0))
    raise RuntimeError(f"Failed to fetch {url} after {_FETCH_MAX_ATTEMPTS} attempts") from last_exc


def build_bundled_task_source(
    *,
    repo: str,
    commit: str,
    max_document_chars: int = 200_000,
) -> TaskSource:
    """Build a task source that prefers documents bundled in the dataset row.

    When a record carries ``document_blobs`` (a dataset exported with
    ``--embed-documents``), the documents are materialized straight from the
    base64 payload — no network. Records without embedded blobs fall back to
    fetching from the pinned commit, so the same spec works for either dataset
    shape.
    """
    import base64

    github_fallback = build_github_task_source(repo=repo, commit=commit, max_document_chars=max_document_chars)

    def _factory(record: Any) -> TaskWorkspace:
        import tempfile

        blobs: Mapping[str, str] = getattr(record, "document_blobs", {}) or {}
        if not blobs:
            return github_fallback(record)
        ws = TaskWorkspace(tempfile.mkdtemp(prefix="harvey_lab_"), max_document_chars=max_document_chars)
        for name, payload in blobs.items():
            dest = ws.documents_dir / name
            dest.parent.mkdir(parents=True, exist_ok=True)
            dest.write_bytes(base64.b64decode(payload))
        return ws

    return _factory


def task_source_from_mapping(mapping: Mapping[str, TaskWorkspace]) -> TaskSource:
    """Factory from a prebuilt ``{task_id: TaskWorkspace}`` mapping (tests)."""

    def _factory(record: Any) -> TaskWorkspace:
        task_id = str(getattr(record, "task_id", "") or "")
        if task_id not in mapping:
            raise KeyError(f"No prebuilt workspace for task_id={task_id!r}.")
        return mapping[task_id]

    return _factory
